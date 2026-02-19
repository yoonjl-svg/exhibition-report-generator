"""
일민미술관 전시보고서 Word(.docx) 생성 엔진
- 기존 구글 문서 기반 보고서 형식을 엄격히 재현
- 제목 체계: I. → 1. → 1) → ① ② ③
- 전시 개요: 불릿 리스트 (● / -)
- 페이지 번호: 우측 하단
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile

from styles import (
    setup_document, set_run_font, add_paragraph, add_horizontal_rule,
    add_section_title, add_subsection_title, add_sub2_title,
    add_detail_title, add_bullet_main, add_bullet_sub, add_arrow_note,
    create_table, create_table_left_aligned,
    add_image, add_images_auto, add_images_2col,
    add_page_break, add_page_numbers_right,
    Colors, Fonts, CIRCLED_NUMBERS, ImageSize,
)
from chart_generator import (
    create_visitor_pie_chart,
    create_weekly_visitors_chart,
    create_budget_comparison_chart,
    create_visitor_type_chart,
)


class ExhibitionReportGenerator:
    """전시보고서 생성기"""

    def __init__(self, data):
        self.data = data
        self.doc = Document()
        self.temp_files = []

    def generate(self, output_path):
        """전체 보고서 생성"""
        setup_document(self.doc)
        add_page_numbers_right(self.doc)

        self._create_toc_page()
        add_page_break(self.doc)

        self._section_1_overview()
        # 전시 개요 후 바로 전시 주제와 내용 (페이지 나누기 없이 이어짐)

        self._section_2_theme()
        add_page_break(self.doc)

        self._section_3_composition()
        add_page_break(self.doc)

        self._section_4_results()
        add_page_break(self.doc)

        self._section_5_promotion()
        add_page_break(self.doc)

        self._section_6_evaluation()

        # 보고서 끝 표기
        add_paragraph(self.doc, "")  # 빈 줄
        add_paragraph(
            self.doc, "끝.",
            size=Fonts.BODY, bold=False,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(12), space_after=Pt(0),
            line_spacing=1.15
        )

        self.doc.save(output_path)
        self._cleanup()
        return output_path

    def _cleanup(self):
        for f in self.temp_files:
            try:
                os.remove(f)
            except OSError:
                pass

    # ══════════════════════════════════════════
    # 목차 페이지 (Page 1)
    # ══════════════════════════════════════════

    def _create_toc_page(self):
        """목차 페이지 — 제목 + 수평선 + 목차 항목들 + 포스터 이미지"""
        title = self.data.get("exhibition_title", "전시 제목")

        # 제목
        add_paragraph(
            self.doc, f"전시보고서 - 《{title}》",
            size=Fonts.TOC_TITLE, bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(12), space_after=Pt(4)
        )
        add_horizontal_rule(self.doc)

        # 목차 항목
        toc_items = [
            "I. 전시 개요",
            "II. 전시 주제와 내용",
            "III. 전시 구성",
            "IV. 전시 결과",
            "V. 홍보 방식 및 언론 보도",
            "VI. 평가 및 개선 방안",
        ]
        for item in toc_items:
            add_paragraph(
                self.doc, item,
                size=Fonts.TOC_ITEM, bold=True,
                space_before=Pt(6), space_after=Pt(6),
                line_spacing=1.3
            )
            add_horizontal_rule(self.doc)

        # 포스터 이미지 (있으면)
        poster = self.data.get("poster_image")
        if poster and os.path.exists(poster):
            add_paragraph(self.doc, "", space_before=Pt(10))
            add_image(self.doc, poster, width=ImageSize.POSTER_WIDTH)

    # ══════════════════════════════════════════
    # I. 전시 개요
    # ══════════════════════════════════════════

    def _section_1_overview(self):
        """I. 전시 개요 — 불릿 리스트 형식 (● / -)"""
        add_section_title(self.doc, "I", "전시 개요")

        ov = self.data.get("overview", {})

        # 기본 정보
        if ov.get("title"):
            add_bullet_main(self.doc, "전시 제목", f"《{ov['title']}》")
        if ov.get("period"):
            add_bullet_main(self.doc, "전시 기간", ov["period"])
        if ov.get("artists"):
            artists = ov["artists"]
            if isinstance(artists, list):
                artists = ", ".join(artists)
            add_bullet_main(self.doc, "참여 작가", artists)

        # 기획진
        if ov.get("chief_curator"):
            add_bullet_main(self.doc, "책임기획", ov["chief_curator"])
        if ov.get("curators"):
            add_bullet_main(self.doc, "기획", ov["curators"])
        if ov.get("coordinators"):
            add_bullet_main(self.doc, "진행", ov["coordinators"])
        if ov.get("curatorial_team"):
            add_bullet_main(self.doc, "학예팀", ov["curatorial_team"])
        if ov.get("pr"):
            add_bullet_main(self.doc, "홍보", ov["pr"])
        if ov.get("sponsors"):
            add_bullet_main(self.doc, "후원", ov["sponsors"])

        # 예산
        if ov.get("total_budget"):
            add_bullet_main(self.doc, "총 사용 예산", ov["total_budget"],
                            bold_value=True, underline_value=True)
            # 하위 항목
            if ov.get("budget_breakdown"):
                for item in ov["budget_breakdown"]:
                    add_bullet_sub(self.doc, item)

        # 수입
        if ov.get("total_revenue"):
            add_bullet_main(self.doc, "총수입", ov["total_revenue"])

        # 프로그램
        if ov.get("programs"):
            add_bullet_main(self.doc, "프로그램", ov["programs"])

        # 운영 인력
        if ov.get("staff_count"):
            add_bullet_main(self.doc, "운영 인력", ov["staff_count"])

        # 관객 수
        if ov.get("visitors"):
            add_bullet_main(self.doc, "관객 수", ov["visitors"],
                            bold_value=True, underline_value=True)

        add_paragraph(self.doc, "")  # 빈 줄

    # ══════════════════════════════════════════
    # II. 전시 주제와 내용
    # ══════════════════════════════════════════

    def _section_2_theme(self):
        """II. 전시 주제와 내용 — 에세이 텍스트"""
        add_section_title(self.doc, "II", "전시 주제와 내용")

        theme = self.data.get("theme_text", "")
        if theme:
            paragraphs = theme.split("\n\n")
            for p_text in paragraphs:
                p_text = p_text.strip()
                if p_text:
                    add_paragraph(
                        self.doc, p_text,
                        size=Fonts.BODY,
                        space_after=Pt(6),
                        line_spacing=1.5,
                        first_line_indent=Cm(0.5)
                    )

    # ══════════════════════════════════════════
    # III. 전시 구성
    # ══════════════════════════════════════════

    def _section_3_composition(self):
        """III. 전시 구성"""
        add_section_title(self.doc, "III", "전시 구성")

        self._sub_exhibition_rooms()
        self._sub_related_programs()
        self._sub_staff()
        self._sub_printed_materials()

    def _sub_exhibition_rooms(self):
        """1. 전시 — 전시실별 도면 + 전경 사진"""
        add_subsection_title(self.doc, "1", "전시")

        rooms = self.data.get("rooms", [])
        for i, room in enumerate(rooms):
            room_name = room.get("name", f"{i + 1}전시실")
            artists = room.get("artists", "")

            add_sub2_title(self.doc, i + 1, room_name)

            # ① 참여 작가
            if artists:
                if isinstance(artists, list):
                    artists = ", ".join(artists)
                add_detail_title(self.doc, CIRCLED_NUMBERS[0], "참여 작가")
                add_paragraph(self.doc, artists, size=Fonts.BODY, left_indent=Cm(0.8))

            # ② 도면
            floor_plan = room.get("floor_plan")
            if floor_plan and os.path.exists(floor_plan):
                add_detail_title(self.doc, CIRCLED_NUMBERS[1], "도면")
                add_image(self.doc, floor_plan)

            # ③ 전경 사진
            photos = room.get("photos", [])
            valid = [p for p in photos if os.path.exists(p)]
            if valid:
                idx = 2 if floor_plan and os.path.exists(floor_plan) else 1
                add_detail_title(self.doc, CIRCLED_NUMBERS[idx], "전경 사진")
                add_images_auto(self.doc, valid)

    def _sub_related_programs(self):
        """2. 전시 연계 프로그램"""
        programs = self.data.get("related_programs", [])
        total_count = len(programs)
        total_participants = 0
        for p in programs:
            try:
                val = str(p.get("participants", "0")).replace(",", "").replace("명", "")
                total_participants += int(val)
            except ValueError:
                pass

        suffix = ""
        if total_count > 0:
            suffix = f" - 총 {total_count}개 프로그램 진행"
            if total_participants > 0:
                suffix += f", {total_participants:,}명 참여"

        add_subsection_title(self.doc, "2", "전시 연계 프로그램", suffix=suffix)

        if programs:
            add_sub2_title(self.doc, "1", "프로그램 운영 내역")

            headers = ["구분", "제목", "일자", "참여 인원", "비고"]
            table_data = []
            for prog in programs:
                table_data.append([
                    prog.get("category", ""),
                    prog.get("title", ""),
                    prog.get("date", ""),
                    prog.get("participants", ""),
                    prog.get("note", "")
                ])

            create_table(
                self.doc, len(table_data), 5,
                data=table_data, headers=headers,
                col_widths=[Cm(2.5), Cm(4.5), Cm(3), Cm(2), Cm(3.5)]
            )

        # 프로그램 사진
        prog_photos = self.data.get("program_photos", [])
        valid = [p for p in prog_photos if os.path.exists(p)]
        if valid:
            add_sub2_title(self.doc, "2", "프로그램 운영 사진")
            add_images_auto(self.doc, valid)

    def _sub_staff(self):
        """3. 전시운영인력"""
        add_subsection_title(self.doc, "3", "전시운영인력")

        staff = self.data.get("staff", {})

        if staff.get("main_staff"):
            add_sub2_title(self.doc, "1", "스태프")
            info = staff["main_staff"]
            if isinstance(info, dict):
                if info.get("count"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[0], "인원")
                    add_paragraph(self.doc, info["count"], left_indent=Cm(0.8))
                if info.get("role"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[1], "역할 및 활동 내용")
                    add_paragraph(self.doc, info["role"], left_indent=Cm(0.8))
            else:
                add_paragraph(self.doc, str(info), left_indent=Cm(0.5))

        if staff.get("volunteers"):
            add_sub2_title(self.doc, "2", "봉사단")
            info = staff["volunteers"]
            if isinstance(info, dict):
                if info.get("count"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[0], "인원")
                    add_paragraph(self.doc, info["count"], left_indent=Cm(0.8))
                if info.get("role"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[1], "역할 및 활동 내용")
                    add_paragraph(self.doc, info["role"], left_indent=Cm(0.8))
            else:
                add_paragraph(self.doc, str(info), left_indent=Cm(0.5))

        if staff.get("support_team"):
            add_sub2_title(self.doc, "3", "50+문화시설지원단")
            info = staff["support_team"]
            if isinstance(info, dict):
                if info.get("count"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[0], "인원")
                    add_paragraph(self.doc, info["count"], left_indent=Cm(0.8))
                if info.get("role"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[1], "역할 및 활동 내용")
                    add_paragraph(self.doc, info["role"], left_indent=Cm(0.8))
            else:
                add_paragraph(self.doc, str(info), left_indent=Cm(0.5))

    def _sub_printed_materials(self):
        """4. 인쇄물"""
        add_subsection_title(self.doc, "4", "인쇄물")

        materials = self.data.get("printed_materials", [])
        if materials:
            headers = ["종류", "제작 수량"]
            table_data = [[m.get("type", ""), m.get("quantity", "")] for m in materials]
            create_table(
                self.doc, len(table_data), 2,
                data=table_data, headers=headers,
                col_widths=[Cm(8), Cm(8)]
            )

        # 인쇄물 이미지
        mat_photos = self.data.get("material_photos", [])
        valid = [p for p in mat_photos if os.path.exists(p)]
        if valid:
            add_paragraph(self.doc, "")
            add_images_auto(self.doc, valid)

    # ══════════════════════════════════════════
    # IV. 전시 결과
    # ══════════════════════════════════════════

    def _section_4_results(self):
        """IV. 전시 결과"""
        add_section_title(self.doc, "IV", "전시 결과")

        self._sub_budget()
        self._sub_revenue()
        self._sub_visitor_composition()

    def _sub_budget(self):
        """1. 예산 및 지출"""
        add_subsection_title(self.doc, "1", "예산 및 지출")

        budget = self.data.get("budget", {})

        # ● 지출 총액 (굵은 + 밑줄)
        if budget.get("total_spent"):
            add_bullet_main(self.doc, "지출 총액", budget["total_spent"],
                            bold_value=True, underline_value=True)

        # - 지출 구성
        if budget.get("breakdown_notes"):
            for note in budget["breakdown_notes"]:
                add_bullet_sub(self.doc, note)

        # 계획 대비 집행 요약 표
        summary = budget.get("summary", [])
        if summary:
            add_paragraph(self.doc, "", space_before=Pt(6))
            headers = ["사업", "계획 예산(원)", "집행 예산(원)", "계획 대비 집행"]
            table_data = []
            for item in summary:
                table_data.append([
                    item.get("category", ""),
                    item.get("planned", ""),
                    item.get("actual", ""),
                    item.get("note", "")
                ])
            create_table(
                self.doc, len(table_data), 4,
                data=table_data, headers=headers,
                col_widths=[Cm(3.5), Cm(4), Cm(4), Cm(4)]
            )

        # → 화살표 주석
        if budget.get("arrow_notes"):
            for note in budget["arrow_notes"]:
                add_arrow_note(self.doc, note)

        # 예산 비교 차트 (자동 생성)
        chart_data = budget.get("chart_data", {})
        if chart_data:
            categories = list(chart_data.keys())
            planned = [chart_data[c].get("planned", 0) for c in categories]
            actual = [chart_data[c].get("actual", 0) for c in categories]
            chart_path = create_budget_comparison_chart(categories, planned, actual)
            self.temp_files.append(chart_path)
            add_image(self.doc, chart_path, is_chart=True)

        # 상세 예산 집행 내역
        details = budget.get("details", [])
        if details:
            add_paragraph(self.doc, "", space_before=Pt(8))
            add_paragraph(self.doc, "예산 집행 내역", size=Fonts.BODY, bold=True,
                          space_after=Pt(4))
            headers = ["사업", "세목", "내역", "금액(원)", "비고"]
            table_data = []
            for item in details:
                table_data.append([
                    item.get("category", ""),
                    item.get("subcategory", item.get("item", "")),
                    item.get("detail", ""),
                    item.get("amount", ""),
                    item.get("note", "")
                ])
            create_table(
                self.doc, len(table_data), 5,
                data=table_data, headers=headers,
                col_widths=[Cm(2.5), Cm(3), Cm(4), Cm(3.5), Cm(3)]
            )

    def _sub_revenue(self):
        """2. 총 관객 수 및 수익 결산"""
        add_subsection_title(self.doc, "2", "총 관객 수 및 수익 결산")

        rev = self.data.get("revenue", {})

        # 1) 총 관객 수
        if rev.get("total_visitors"):
            add_sub2_title(self.doc, "1", f"총 관객 수 {rev['total_visitors']}")

            if rev.get("daily_average"):
                add_bullet_main(self.doc, "일평균 관객", rev["daily_average"])
            if rev.get("visitor_notes"):
                for note in rev["visitor_notes"]:
                    add_bullet_sub(self.doc, note)

        # 2) 총 수입
        if rev.get("total_revenue"):
            add_sub2_title(self.doc, "2", f"총 수입 {rev['total_revenue']}")

            if rev.get("ticket_revenue"):
                add_bullet_main(self.doc, "입장 수입", rev["ticket_revenue"])
            if rev.get("partnership_revenue"):
                add_bullet_main(self.doc, "제휴 수입", rev["partnership_revenue"])
            if rev.get("revenue_notes"):
                for note in rev["revenue_notes"]:
                    add_bullet_sub(self.doc, note)

    def _sub_visitor_composition(self):
        """3. 관객 구성"""
        add_subsection_title(self.doc, "3", "관객 구성")

        vc = self.data.get("visitor_composition", {})

        # 주석
        if vc.get("note"):
            add_paragraph(self.doc, f"※ {vc['note']}", size=Fonts.BODY, bold=True,
                          space_after=Pt(6))

        # 입장권별 파이차트
        ticket_type = vc.get("ticket_type", {})
        if ticket_type:
            chart_path = create_visitor_pie_chart(ticket_type, title="입장권별 관객 구성")
            self.temp_files.append(chart_path)
            add_image(self.doc, chart_path, is_chart=True)

        # 분석 불릿
        if vc.get("ticket_analysis"):
            for item in vc["ticket_analysis"]:
                if item.startswith("→"):
                    add_arrow_note(self.doc, item[1:].strip())
                elif item.startswith("-"):
                    add_bullet_sub(self.doc, item[1:].strip())
                else:
                    add_bullet_main(self.doc, None, item, bold_value=True,
                                    underline_value=True)

        # 유형별 관객 수 파이차트
        visitor_type = vc.get("visitor_type", {})
        if visitor_type:
            chart_path = create_visitor_type_chart(visitor_type)
            self.temp_files.append(chart_path)
            add_paragraph(self.doc, "", space_before=Pt(8))
            add_image(self.doc, chart_path, is_chart=True)

        # 주별 관객 수 바 차트
        weekly = vc.get("weekly_visitors", {})
        if weekly:
            chart_path = create_weekly_visitors_chart(weekly)
            self.temp_files.append(chart_path)
            add_paragraph(self.doc, "", space_before=Pt(8))
            add_image(self.doc, chart_path, is_chart=True)

        # 관객 분석 텍스트
        analysis = vc.get("analysis", "")
        if analysis:
            add_paragraph(self.doc, analysis, size=Fonts.BODY,
                          space_before=Pt(8), line_spacing=1.5)

    # ══════════════════════════════════════════
    # V. 홍보 방식 및 언론 보도
    # ══════════════════════════════════════════

    def _section_5_promotion(self):
        """V. 홍보 방식 및 언론 보도"""
        add_section_title(self.doc, "V", "홍보 방식 및 언론 보도")

        self._sub_promo_methods()
        self._sub_press_list()
        self._sub_membership()
        self._sub_ima_on()

    def _sub_promo_methods(self):
        """1. 홍보 방식"""
        add_subsection_title(self.doc, "1", "홍보 방식")

        promo = self.data.get("promotion", {})
        categories = [
            ("advertising", "광고"),
            ("press_release", "보도자료"),
            ("web_invitation", "웹 초청장"),
            ("newsletter", "뉴스레터"),
            ("sns", "SNS"),
            ("other", "그 외"),
        ]

        num = 1
        for key, label in categories:
            content = promo.get(key, "")
            if content:
                add_sub2_title(self.doc, num, label)
                # 여러 줄이면 각각 불릿으로
                lines = content.split("\n")
                for line in lines:
                    line = line.strip()
                    if line:
                        add_bullet_main(self.doc, None, line)
                num += 1

        # 홍보 이미지
        promo_photos = self.data.get("promotion_photos", [])
        valid = [p for p in promo_photos if os.path.exists(p)]
        if valid:
            add_images_auto(self.doc, valid)

    def _sub_press_list(self):
        """2. 언론보도 리스트"""
        add_subsection_title(self.doc, "2", "언론보도 리스트")

        press = self.data.get("press_coverage", {})

        # 일간지 및 월간지
        print_media = press.get("print_media", [])
        if print_media:
            add_sub2_title(self.doc, "1", "일간지 및 월간지")
            headers = ["매체명", "일자", "제목", "비고"]
            table_data = [[
                item.get("outlet", ""),
                item.get("date", ""),
                item.get("title", ""),
                item.get("note", "")
            ] for item in print_media]
            create_table(
                self.doc, len(table_data), 4,
                data=table_data, headers=headers,
                col_widths=[Cm(3), Cm(3), Cm(6), Cm(4)]
            )

        # 온라인 매체
        online_media = press.get("online_media", [])
        if online_media:
            add_sub2_title(self.doc, "2", "온라인 매체")
            headers = ["매체명", "일자", "제목", "URL"]
            table_data = [[
                item.get("outlet", ""),
                item.get("date", ""),
                item.get("title", ""),
                item.get("url", "")
            ] for item in online_media]
            create_table(
                self.doc, len(table_data), 4,
                data=table_data, headers=headers,
                col_widths=[Cm(3), Cm(3), Cm(5), Cm(5)]
            )

    def _sub_membership(self):
        """3. 멤버십 커뮤니케이션"""
        membership = self.data.get("membership", "")
        if membership:
            add_subsection_title(self.doc, "3", "멤버십 커뮤니케이션")
            add_paragraph(self.doc, membership, size=Fonts.BODY, line_spacing=1.4)

    def _sub_ima_on(self):
        """4. IMA ON"""
        ima_on = self.data.get("ima_on", {})
        if ima_on and (ima_on.get("content") or ima_on.get("photos")):
            title = ima_on.get("title", "IMA ON")
            add_subsection_title(self.doc, "4", title)

            content = ima_on.get("content", "")
            if content:
                add_paragraph(self.doc, content, size=Fonts.BODY, line_spacing=1.4)

            photos = ima_on.get("photos", [])
            valid = [p for p in photos if os.path.exists(p)]
            if valid:
                add_images_auto(self.doc, valid)

    # ══════════════════════════════════════════
    # VI. 평가 및 개선 방안
    # ══════════════════════════════════════════

    def _section_6_evaluation(self):
        """VI. 평가 및 개선 방안"""
        add_section_title(self.doc, "VI", "평가 및 개선 방안")

        self._sub_evaluation()

    def _sub_evaluation(self):
        """1. 평가 — 긍정 평가 + 긍정 후기 표 → 부정 평가 + 부정 후기 표 → 개선 방안"""
        add_subsection_title(self.doc, "1", "평가")

        evaluation = self.data.get("evaluation", {})
        reviews = self.data.get("visitor_reviews", [])

        # 후기를 긍정/부정으로 분류
        positive_reviews = [r for r in reviews if r.get("category", "").strip() in ("긍정", "긍정적")]
        negative_reviews = [r for r in reviews if r.get("category", "").strip() in ("부정", "부정적", "건의", "불만")]

        sub_num = 1

        # 1) 긍정 평가
        positive = evaluation.get("positive", [])
        if positive or positive_reviews:
            add_sub2_title(self.doc, sub_num, "긍정 평가")
            for item in positive:
                add_bullet_main(self.doc, None, item)

            # 긍정 후기 표
            if positive_reviews:
                add_paragraph(self.doc, "", space_before=Pt(4))
                headers = ["분류", "상세 내용(인용)", "출처"]
                table_data = [[
                    r.get("category", "긍정"),
                    r.get("content", ""),
                    r.get("source", "")
                ] for r in positive_reviews]
                create_table(
                    self.doc, len(table_data), 3,
                    data=table_data, headers=headers,
                    col_widths=[Cm(2.5), Cm(9.5), Cm(3)]
                )
            sub_num += 1

        # 2) 부정 평가
        negative = evaluation.get("negative", [])
        if negative or negative_reviews:
            add_sub2_title(self.doc, sub_num, "부정 평가")
            for item in negative:
                add_bullet_main(self.doc, None, item)

            # 부정 후기 표
            if negative_reviews:
                add_paragraph(self.doc, "", space_before=Pt(4))
                headers = ["분류", "상세 내용(인용)", "출처"]
                table_data = [[
                    r.get("category", "부정"),
                    r.get("content", ""),
                    r.get("source", "")
                ] for r in negative_reviews]
                create_table(
                    self.doc, len(table_data), 3,
                    data=table_data, headers=headers,
                    col_widths=[Cm(2.5), Cm(9.5), Cm(3)]
                )
            sub_num += 1

        # 3) 개선 방안
        improvements = evaluation.get("improvements", [])
        if improvements:
            add_sub2_title(self.doc, sub_num, "개선 방안")
            for item in improvements:
                add_bullet_main(self.doc, None, item)


# ──────────────────────────────────────────────
# 편의 함수
# ──────────────────────────────────────────────

def generate_report(data, output_path):
    """보고서 생성"""
    generator = ExhibitionReportGenerator(data)
    return generator.generate(output_path)


# ──────────────────────────────────────────────
# 테스트
# ──────────────────────────────────────────────

if __name__ == "__main__":
    sample_data = {
        "exhibition_title": "하이퍼 옐로우",
        "exhibition_period": "2024.09.06 - 2024.10.27",
        "overview": {
            "title": "하이퍼 옐로우",
            "period": "2024.09.06(금) - 2024.10.27(일) (52일간, 월요일 휴관)",
            "artists": ["구정연", "이미래", "장서영"],
            "chief_curator": "최지현",
            "curators": "윤지현",
            "coordinators": "윤율리",
            "pr": "전다희",
            "sponsors": "한국문화예술위원회",
            "total_budget": "약 1억 4천 2백만 원(142,438,012원)",
            "budget_breakdown": [
                "지출 구성: 전시비 130,773,012원 / 부대비 11,665,000원"
            ],
            "total_revenue": "49,574,000원",
            "programs": "총 8개(27회) 프로그램 진행, 719명 참여",
            "staff_count": "스태프 10명, 봉사단 12명",
            "visitors": "7,009명"
        },
        "theme_text": "《하이퍼 옐로우》는 노란색이라는 색채를 통해 동시대 미술의 감각적 경험을 탐구하는 전시이다. 세 명의 작가가 각자의 매체와 언어로 '옐로우'라는 주제에 접근하며, 관객에게 시각적, 촉각적, 청각적 경험을 총체적으로 제공한다.\n\n전시는 회화, 영상, 설치 등 다양한 매체를 아우르며, 색채가 단순한 시각적 요소를 넘어 감정과 사회적 의미를 어떻게 전달하는지를 탐구한다. 작가들은 각자의 방식으로 노란색의 다층적 의미를 풀어내며, 관객이 능동적으로 작품을 경험할 수 있는 공간을 만들어낸다.",
        "rooms": [
            {"name": "1전시실", "artists": ["구정연", "이미래"]},
            {"name": "2전시실", "artists": ["장서영"]},
        ],
        "related_programs": [
            {"category": "아티스트 토크", "title": "작가와의 대화", "date": "2024.09.20", "participants": "50", "note": ""},
            {"category": "큐레이터 투어", "title": "전시 해설", "date": "2024.09.14", "participants": "30", "note": ""},
            {"category": "워크숍", "title": "색채 워크숍", "date": "2024.10.05", "participants": "25", "note": "가족 대상"},
        ],
        "program_photos": [],
        "staff": {
            "main_staff": {"count": "총 10명", "role": "전시 안내, 작품 모니터링, 관객 응대"},
            "volunteers": {"count": "총 12명 (제17기)", "role": "전시 안내 보조, 교육 프로그램 지원"},
        },
        "printed_materials": [
            {"type": "리플렛", "quantity": "5,000부"},
            {"type": "포스터(B1)", "quantity": "200부"},
            {"type": "초청장", "quantity": "300부"},
        ],
        "material_photos": [],
        "budget": {
            "total_spent": "약 1억 4천 2백만 원(142,438,012원)",
            "breakdown_notes": [
                "지출 구성: 전시비 130,773,012원 / 부대비 11,665,000원"
            ],
            "summary": [
                {"category": "전시 사업비", "planned": "125,200,000", "actual": "130,773,012", "note": "104.4%"},
                {"category": "부대 사업비", "planned": "12,000,000", "actual": "11,665,000", "note": "97.2%"},
            ],
            "arrow_notes": [
                "전시 예산의 104.2% 사용: 작가 설치비 추가 지출"
            ],
            "chart_data": {
                "전시 사업비": {"planned": 125200000, "actual": 130773012},
                "부대 사업비": {"planned": 12000000, "actual": 11665000},
            },
            "details": [
                {"category": "전시비", "subcategory": "작품 제작비", "detail": "작가 3인 제작 지원", "amount": "45,000,000", "note": ""},
                {"category": "전시비", "subcategory": "설치비", "detail": "전시장 설치/철거", "amount": "28,000,000", "note": ""},
                {"category": "전시비", "subcategory": "디자인비", "detail": "인쇄물 디자인", "amount": "8,500,000", "note": ""},
                {"category": "부대비", "subcategory": "교육 프로그램", "detail": "워크숍 3회", "amount": "5,000,000", "note": ""},
            ],
        },
        "revenue": {
            "total_visitors": "7,009명",
            "daily_average": "135명",
            "visitor_notes": ["짧은 전시 기간(52일) 대비 양호한 관객 수 기록"],
            "total_revenue": "49,574,000원",
            "ticket_revenue": "42,574,000원",
            "partnership_revenue": "7,000,000원",
        },
        "visitor_composition": {
            "note": "티켓 권종 기준으로 작성",
            "ticket_type": {"일반": 3500, "학생": 800, "초대권": 500, "예술인패스": 509, "기타 할인": 1700},
            "ticket_analysis": [
                "예술인패스 관객이 전체의 7.3%를 차지하며 높은 비율 기록",
                "→ 예술인패스 관객 대상 특화 프로그램 기획 검토 필요",
                "- 학생 관객 비율 11.4% 기록",
            ],
            "visitor_type": {"개인": 5500, "미술대학 단체": 600, "기타 단체": 400, "오프닝 리셉션": 509},
            "weekly_visitors": {"1주": 1200, "2주": 980, "3주": 850, "4주": 900, "5주": 1050, "6주": 800, "7주": 1229},
            "analysis": "개인 관객이 전체의 약 78%를 차지하며, 미술대학 단체 방문이 전시 후반부에 집중되었다. 주별 관객 수는 개막 첫 주와 마지막 주에 높게 나타났다.",
        },
        "promotion": {
            "advertising": "서울 주요 지하철역(광화문, 을지로, 종각) 포스터 게시\n미술전문 주간지 광고 게재 (월간미술, 아트인컬쳐)",
            "press_release": "보도자료 3회 배포 (개막 전, 개막 당일, 연계 프로그램)",
            "sns": "인스타그램 게시물 35회, 릴스 8회\n페이스북 게시물 15회",
        },
        "promotion_photos": [],
        "press_coverage": {
            "print_media": [
                {"outlet": "조선일보", "date": "2024.09.07", "title": "일민미술관, 하이퍼 옐로우 전시 개막", "note": ""},
                {"outlet": "경향신문", "date": "2024.09.10", "title": "색채로 탐구하는 동시대 미술", "note": ""},
            ],
            "online_media": [
                {"outlet": "아트인사이트", "date": "2024.09.08", "title": "하이퍼 옐로우 리뷰", "url": ""},
                {"outlet": "네이버 아트", "date": "2024.09.12", "title": "이번 주 추천 전시", "url": ""},
            ]
        },
        "membership": "멤버십 회원 대상 프리뷰 행사 개최 (2024.09.05)\n참여 인원: 85명\n멤버십 전용 뉴스레터 발송 3회",
        "ima_on": {
            "title": "IMA Critics",
            "content": "비평가 초청 리뷰 프로그램 운영\n참여 비평가: 2명\n온라인 게재: 일민미술관 웹사이트 및 SNS",
        },
        "evaluation": {
            "positive": [
                "20~30대 젊은 관객층의 높은 호응도 (전체의 66%)",
                "SNS 바이럴을 통한 자연 유입 관객 증가 (인스타그램 릴스 조회수 합계 12만회)",
                "언론 보도 총 15건 달성",
                "연계 프로그램 참여율 높음 (평균 정원 대비 95% 참여)",
            ],
            "negative": [
                "평일 관객 유입 저조 (주말 대비 평일 관객 40% 수준)",
                "전시 안내 동선이 불명확하다는 피드백 다수",
            ],
            "improvements": [
                "평일 관객 유입 확대 필요 (주말 대비 평일 관객 40% 수준)",
                "외국인 관객 유치를 위한 다국어 안내 확대",
                "전시 기간 중 중간 점검을 통한 홍보 전략 조정 필요",
            ]
        },
        "visitor_reviews": [
            {"category": "긍정", "content": "색채를 통한 감각적 경험이 인상적이었습니다.", "source": "방명록"},
            {"category": "긍정", "content": "작가와의 대화 프로그램에서 깊은 이야기를 들을 수 있었어요.", "source": "SNS"},
            {"category": "부정", "content": "안내 리플렛에 영문 설명이 없어 아쉬웠습니다.", "source": "방명록"},
            {"category": "부정", "content": "전시장 내 동선 안내가 부족했습니다.", "source": "SNS"},
        ]
    }

    output = "/tmp/test_report_v2.docx"
    generate_report(sample_data, output)
    print(f"보고서 생성 완료: {output}")
